from docx import Document
import random

def create_docx(summary_text: str, date_range: str) -> str:
    doc = Document()
    doc.add_heading("Haftalık Rapor", 0)
    doc.add_paragraph(f"Tarih: {date_range}")
    doc.add_paragraph(summary_text)
    filename = f"Haftalik_Rapor_{random.randint(10000, 99999)}.docx"
    doc.save(filename)
    return filename
